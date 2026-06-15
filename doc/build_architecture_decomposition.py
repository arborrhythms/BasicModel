#!/usr/bin/env python3
"""Build the BasicModel architecture decomposition document.

Generates:
  - doc/architecture_decomposition_assets/*.png
  - doc/BasicModel_Architecture_Decomposition.docx

The diagrams are intentionally drawn as plain PNGs so the DOCX/PDF render
path does not depend on Mermaid, Graphviz, or browser-side script execution.
"""

from __future__ import annotations

import ast
import math
import textwrap
from pathlib import Path
from xml.sax.saxutils import escape

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as PdfImage,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)


ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "doc"
BIN_DIR = ROOT / "bin"
ASSET_DIR = DOC_DIR / "architecture_decomposition_assets"
DOCX_PATH = DOC_DIR / "BasicModel_Architecture_Decomposition.docx"
PDF_PATH = DOC_DIR / "BasicModel_Architecture_Decomposition.pdf"

BLUE = "#2E74B5"
DARK_BLUE = "#1F4D78"
INK = "#1B2430"
MUTED = "#5E6A75"
LINE = "#B8C3CF"
FILL = "#F4F7FA"
FILL_BLUE = "#E8EEF5"
FILL_GREEN = "#EAF5EF"
FILL_GOLD = "#FFF4D6"
FILL_RED = "#FCE8E6"
WHITE = "#FFFFFF"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Helvetica.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


FONT_TITLE = font(42, True)
FONT_H = font(40, True)
FONT = font(30)
FONT_SMALL = font(27)
FONT_TINY = font(22)


def text_size(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont) -> tuple[int, int]:
    bbox = draw.textbbox((0, 0), text, font=fnt)
    return bbox[2] - bbox[0], bbox[3] - bbox[1]


def wrap_lines(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for raw in text.split("\n"):
        words = raw.split()
        if not words:
            lines.append("")
            continue
        current = words[0]
        for word in words[1:]:
            test = f"{current} {word}"
            if text_size(draw, test, fnt)[0] <= max_width:
                current = test
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return lines


def draw_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    title: str,
    body: str = "",
    fill: str = FILL,
    outline: str = LINE,
) -> None:
    x0, y0, x1, y1 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    inner = x1 - x0 - 28
    y = y0 + 16
    for line in wrap_lines(draw, title, FONT_H, inner):
        w, h = text_size(draw, line, FONT_H)
        draw.text((x0 + (x1 - x0 - w) / 2, y), line, fill=INK, font=FONT_H)
        y += h + 4
    if body:
        y += 5
        for line in wrap_lines(draw, body, FONT_SMALL, inner):
            w, h = text_size(draw, line, FONT_SMALL)
            draw.text((x0 + (x1 - x0 - w) / 2, y), line, fill=MUTED, font=FONT_SMALL)
            y += h + 4


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#667789", width: int = 3) -> None:
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 12
    p1 = (end[0] - size * math.cos(angle - math.pi / 6), end[1] - size * math.sin(angle - math.pi / 6))
    p2 = (end[0] - size * math.cos(angle + math.pi / 6), end[1] - size * math.sin(angle + math.pi / 6))
    draw.polygon([end, p1, p2], fill=color)


def poly_arrow(draw: ImageDraw.ImageDraw, points: list[tuple[int, int]], color: str = "#667789", width: int = 3) -> None:
    draw.line(points, fill=color, width=width, joint="curve")
    arrow(draw, points[-2], points[-1], color=color, width=width)


def label(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, fnt=FONT_TINY, color=MUTED) -> None:
    x, y = xy
    for line in text.split("\n"):
        draw.text((x, y), line, font=fnt, fill=color)
        y += text_size(draw, line, fnt)[1] + 4


def canvas(name: str, width: int = 2400, height: int = 1600) -> tuple[Image.Image, ImageDraw.ImageDraw, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (width, height), WHITE)
    draw = ImageDraw.Draw(img)
    return img, draw, ASSET_DIR / name


def save(img: Image.Image, path: Path) -> None:
    img.save(path, "PNG", optimize=True)


def static_architecture_diagram() -> Path:
    img, d, path = canvas("static_architecture.png", 2400, 1700)
    d.text((70, 45), "BasicModel Static Architecture: 24 UML Blocks", fill=INK, font=FONT_TITLE)
    label(d, (72, 95), "Grouped blocks enumerate the concrete Model, Space, Layer, grammar, truth, and taxonomy surfaces.")

    boxes = {
        "train": (70, 165, 390, 285, "1. Train CLI", "train.py\nlocal/remote", FILL),
        "serve": (70, 320, 390, 440, "2. Serve API", "serve.py\nchat endpoint", FILL),
        "config": (70, 475, 390, 595, "3. XML Config", "architecture\ntraining data", FILL),
        "data": (70, 630, 390, 750, "4. Data", "TheData\nSentenceStream", FILL),

        "models": (470, 245, 820, 395, "5. Models", "BaseModel\nBasicModel\nModelFactory", FILL_BLUE),
        "loss": (470, 435, 820, 565, "24. Runtime", "runBatch/runEpoch\nloss + optimizer", FILL_BLUE),
        "input": (470, 660, 820, 790, "6. InputSpace", "raw input -> dual view", FILL_BLUE),

        "ps": (905, 170, 1265, 300, "7. PartSpace", "PS extent tower\nbottom-up synthesis", FILL_GREEN),
        "sigma": (905, 330, 1265, 450, "8. Sigma Fold", "synthesis / union", FILL_GREEN),
        "meron": (905, 480, 1265, 600, "9. Meronymic", "router + analyzer\ntowers", FILL_GREEN),
        "pscb": (905, 630, 1265, 750, "10. PS Codebook", "atoms, lexicon\nextension learning", FILL_GREEN),

        "cs": (1345, 150, 1705, 280, "11. ConceptualSpace", "STM host\ncallosum mixing", FILL_GOLD),
        "cc": (1345, 310, 1705, 430, "12. Corpus Callosum", "nameless factored\n2N -> N mixing", FILL_GOLD),
        "symtab": (1345, 460, 1705, 580, "13. CS Symbol Table", "reference rows\nwide/deep remap", FILL_GOLD),
        "stm": (1345, 610, 1705, 730, "14. STM", "working set\nsentence root", FILL_GOLD),

        "word": (1785, 130, 2265, 255, "15. WordSpace", "grammar CPU\ntruth/discourse host", FILL_BLUE),
        "attn": (1785, 285, 2265, 405, "16. Semantic Attention", "single intent primes\nall codebooks", FILL_BLUE),
        "router": (1785, 435, 2265, 555, "17. LanguageLayer", "signal router\ncopy/reduce", FILL_BLUE),
        "grammar": (1785, 585, 2265, 705, "18. Grammar", "rules, signatures\nprobabilities", FILL_BLUE),
        "gops": (1785, 735, 2265, 855, "19. Grammar Ops", "logic, mereology\nlift/lower", FILL_BLUE),
        "truth": (1785, 885, 2265, 1010, "20. TruthLayer", "consistency gate\npropositions", FILL_RED),
        "tax": (1785, 1040, 2265, 1165, "21. Taxonomy", "RelativeTruth\nintensional hierarchy", FILL_RED),
        "ws": (1345, 925, 1705, 1055, "22. WholeSpace", "SS intent tower\ntop-down analysis", FILL_GREEN),
        "pi": (1345, 1085, 1705, 1205, "23. Pi Fold", "analysis / intersection", FILL_GREEN),
        "output": (905, 1005, 1265, 1135, "24. OutputSpace", "prediction head", FILL_BLUE),
    }

    for item in boxes.values():
        draw_box(d, item[:4], item[4], item[5], item[6])

    arrow(d, (390, 535), (470, 320))
    arrow(d, (390, 225), (470, 320))
    arrow(d, (390, 380), (470, 320))
    arrow(d, (390, 690), (470, 725))
    arrow(d, (645, 395), (645, 435))
    arrow(d, (820, 725), (905, 235))
    arrow(d, (1085, 300), (1085, 330))
    arrow(d, (1085, 450), (1085, 480))
    arrow(d, (1265, 235), (1345, 215))
    arrow(d, (1525, 280), (1525, 310))
    arrow(d, (1525, 430), (1525, 460))
    arrow(d, (1525, 580), (1525, 610))
    arrow(d, (1705, 215), (1785, 195))
    arrow(d, (2025, 255), (2025, 285))
    arrow(d, (2025, 405), (2025, 435))
    arrow(d, (2025, 555), (2025, 585))
    arrow(d, (2025, 705), (2025, 735))
    arrow(d, (2025, 855), (2025, 885))
    arrow(d, (2025, 1010), (2025, 1040))
    arrow(d, (1785, 1100), (1705, 990))
    arrow(d, (1525, 1055), (1525, 1085))
    arrow(d, (1345, 1150), (1265, 1070))
    arrow(d, (905, 1070), (820, 500))
    arrow(d, (1785, 345), (1265, 690))
    arrow(d, (1785, 345), (1705, 990))

    label(d, (74, 1560), "Extension learning: PS/SS codebooks learn word extent from use. Intension learning: TruthLayer admits consistent propositions into Taxonomy and symbol geometry.")
    save(img, path)
    return path


def dynamic_loops_diagram() -> Path:
    img, d, path = canvas("dynamic_loops.png", 2400, 1500)
    d.text((70, 45), "Dynamic Algorithms: Subsymbolic Order, Symbolic Order, Callosum, Discourse", fill=INK, font=FONT_TITLE)

    draw_box(d, (90, 170, 440, 310), "InputSpace", "dual view:\natoms + unity", FILL)
    draw_box(d, (535, 130, 2220, 470), "A. Subsymbolic order loop", "<subsymbolicOrder>: repeat PS -> CS -> WS refinement; parallel parse produces current intent.", FILL_BLUE)
    draw_box(d, (620, 260, 900, 405), "PS", "Sigma synthesis\nextent tower", FILL_GREEN)
    draw_box(d, (990, 260, 1270, 405), "CS", "callosum mixing\nSTM write", FILL_GOLD)
    draw_box(d, (1360, 260, 1640, 405), "WS", "Pi analysis\nintent tower", FILL_GREEN)
    draw_box(d, (1730, 260, 2110, 405), "Intent", "semantic attention\nprimes both towers", FILL_BLUE)

    arrow(d, (440, 240), (620, 330))
    arrow(d, (900, 330), (990, 330))
    arrow(d, (1270, 330), (1360, 330))
    arrow(d, (1640, 330), (1730, 330))
    arrow(d, (1915, 405), (760, 405), color="#4D8F68")
    label(d, (1060, 415), "loopback: attention biases recognition/retrieval, not rule dispatch", FONT_TINY)

    draw_box(d, (535, 560, 2220, 900), "B. Symbolic order loop", "<symbolicOrder>: serial/grammatical processing over words in STM.", FILL_BLUE)
    draw_box(d, (620, 690, 900, 830), "STM push", "word/root\nenters memory", FILL_GOLD)
    draw_box(d, (990, 690, 1270, 830), "LanguageLayer", "score copy/reduce\nroutes", FILL_BLUE)
    draw_box(d, (1360, 690, 1640, 830), "Ops", "grammar +\nmereology", FILL_BLUE)
    draw_box(d, (1730, 690, 2110, 830), "Writes", "C-side read ops\nS-side codebook ops", FILL_GREEN)
    arrow(d, (900, 760), (990, 760))
    arrow(d, (1270, 760), (1360, 760))
    arrow(d, (1640, 760), (1730, 760))
    arrow(d, (1920, 865), (760, 865), color="#667789")

    draw_box(d, (535, 990, 2220, 1320), "C. Inter-sentence prediction", "InterSentenceLayer predicts the next STM end-state/root and stages a C-prior.", FILL_GOLD)
    draw_box(d, (620, 1110, 900, 1250), "End-state", "STM root\npayload", FILL_GOLD)
    draw_box(d, (990, 1110, 1270, 1250), "LTM chain", "recent sentence\nend-states", FILL_GOLD)
    draw_box(d, (1360, 1110, 1640, 1250), "Predict", "next depth\n+ payload", FILL_GOLD)
    draw_box(d, (1730, 1110, 2110, 1250), "C-prior + loss", "stage prior\naccumulate L_inter", FILL_GOLD)
    arrow(d, (900, 1180), (990, 1180))
    arrow(d, (1270, 1180), (1360, 1180))
    arrow(d, (1640, 1180), (1730, 1180))

    save(img, path)
    return path


def truth_taxonomy_diagram() -> Path:
    img, d, path = canvas("truth_taxonomy.png", 2400, 1450)
    d.text((70, 45), "Truth, Intension, and Taxonomic Encoding", fill=INK, font=FONT_TITLE)

    draw_box(d, (90, 170, 520, 325), "Grammar parse", "candidate proposition\nfrom STM and grammar ops", FILL_BLUE)
    draw_box(d, (650, 170, 1080, 325), "TruthLayer", "ascertain consistency:\nsupport / conflict / ignorance", FILL_RED)
    draw_box(d, (1210, 85, 1660, 245), "AbsoluteTruth", "admit fact as symbolic\ntruth activation", FILL_RED)
    draw_box(d, (1210, 330, 1660, 490), "RelativeTruth", "relation over propositions:\npart/equal/query", FILL_RED)
    draw_box(d, (1790, 85, 2250, 245), "Symbol codebook", "geometric storage of\ntruth-bearing symbols", FILL_GREEN)
    draw_box(d, (1790, 330, 2250, 490), "Taxonomy", "intensional hierarchy:\nman < mortal", FILL_GREEN)
    draw_box(d, (1210, 575, 1660, 735), "Clarify / reject", "do not encode inconsistent\nor underdetermined claims", FILL)
    draw_box(d, (1790, 575, 2250, 735), "Token promotion", "tokens seek nearest\ncontaining symbolic whole", FILL_GOLD)

    arrow(d, (520, 247), (650, 247))
    arrow(d, (1080, 225), (1210, 165))
    arrow(d, (1080, 270), (1210, 410))
    arrow(d, (1080, 300), (1210, 650))
    arrow(d, (1660, 165), (1790, 165))
    arrow(d, (1660, 410), (1790, 410))
    arrow(d, (2020, 490), (2020, 575))

    draw_box(d, (90, 890, 1080, 1235), "Extension learning", "Subsymbolic/codebook learning discovers the extension of words from use: tokens, contexts, percepts, and nearest-codebook recognition.", FILL_BLUE)
    draw_box(d, (1210, 890, 2250, 1235), "Intension learning", "Consistent relative truths define words through symbolic containment. If man < mortal, the symbol for man is part of the symbol for mortal; Socrates can be promoted to the nearest containing type when evidence supports it.", FILL_GREEN)
    arrow(d, (1080, 1060), (1210, 1060), color="#667789")

    label(d, (90, 1320), "Key distinction: grammar learns how propositions are formed; TruthLayer and Taxonomy decide which propositions become admissible structure.")
    save(img, path)
    return path


def extract_classes(path: Path) -> list[str]:
    tree = ast.parse(path.read_text())
    return [node.name for node in tree.body if isinstance(node, ast.ClassDef)]


def class_inventory() -> dict[str, list[str]]:
    return {
        "Models.py": extract_classes(BIN_DIR / "Models.py"),
        "Spaces.py": extract_classes(BIN_DIR / "Spaces.py"),
        "Layers.py": extract_classes(BIN_DIR / "Layers.py"),
        "Language.py grammar/word classes": [
            name
            for name in extract_classes(BIN_DIR / "Language.py")
            if name.endswith("Layer")
            or name in {"Grammar", "RuleCodebook", "LanguageLayer", "Taxonomy", "ObjectSubSpace", "WordSubSpace", "NeuralToolUser"}
        ],
        "perceptual_analyzer.py": extract_classes(BIN_DIR / "perceptual_analyzer.py"),
    }


def set_font(run, name: str = "Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)


def set_style_font(style, name: str, size: int, color: str | None = None, bold: bool | None = None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    if bold is not None:
        style.font.bold = bold
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)


def add_page_break_before(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    page_break = OxmlElement("w:pageBreakBefore")
    p_pr.append(page_break)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style_font(styles["Normal"], "Calibri", 11, "#1B2430")
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.25
    set_style_font(styles["Heading 1"], "Calibri", 16, BLUE, True)
    styles["Heading 1"].paragraph_format.space_before = Pt(18)
    styles["Heading 1"].paragraph_format.space_after = Pt(10)
    set_style_font(styles["Heading 2"], "Calibri", 13, BLUE, True)
    styles["Heading 2"].paragraph_format.space_before = Pt(14)
    styles["Heading 2"].paragraph_format.space_after = Pt(7)
    set_style_font(styles["Heading 3"], "Calibri", 12, DARK_BLUE, True)
    styles["Heading 3"].paragraph_format.space_before = Pt(10)
    styles["Heading 3"].paragraph_format.space_after = Pt(5)
    set_style_font(styles["Title"], "Calibri", 24, "#0B2545", True)

    for name, size in (("Caption", 9), ("Intense Quote", 10)):
        if name in styles:
            set_style_font(styles[name], "Calibri", size, "#5E6A75")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("BasicModel architecture decomposition")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("5E6A75")
    set_font(run)
    return doc


def para(doc: Document, text: str = "", style: str | None = None, bold_label: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_label:
        r = p.add_run(bold_label)
        r.bold = True
        set_font(r)
        r2 = p.add_run(text)
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_font(r)


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.5))
    c = doc.add_paragraph(style="Caption")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption)
    r.italic = True
    set_font(r)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("BasicModel Architecture Decomposition")
    set_font(run)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    r = subtitle.add_run(
        "Static modules, meronymic towers, grammar/truth machinery, and dynamic learning loops"
    )
    set_font(r)
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string("5E6A75")


def build_docx(images: dict[str, Path]) -> None:
    doc = setup_document()
    add_title(doc)

    para(
        doc,
        "This brief decomposes BasicModel into roughly two dozen architectural blocks, "
        "then expands those blocks into the concrete Model, Space, Layer, grammar, "
        "truth, codebook, and taxonomy modules present in the codebase. It distinguishes "
        "extension learning over codebooks from intension learning through TruthLayer "
        "and Taxonomy.",
    )

    doc.add_heading("1. Static Architecture", level=1)
    add_figure(
        doc,
        images["static"],
        "Figure 1. Static decomposition: entrypoints, model factory/runtime, five-space body, WordSpace grammar/truth host, codebook towers, and training runtime.",
    )

    p = para(
        doc,
        "InputSpace -> PartSpace -> ConceptualSpace -> WholeSpace -> OutputSpace.",
        bold_label="Pipeline. ",
    )
    for run in p.runs:
        set_font(run)
    para(
        doc,
        "The WordSpace/WordSubSpace sidecar owns the grammar CPU, semantic attention, truth admission, taxonomy, and discourse prediction that act on the STM and codebook towers."
    )

    doc.add_heading("2. Modules and Responsibilities", level=1)
    add_bullets(
        doc,
        [
            "Models: BaseModel owns configuration hydration, training, persistence, and inference scaffolding; BasicModel assembles the space pipeline; ModelFactory validates XML and drives train/eval runs.",
            "InputSpace: converts raw input into the dual view used by the part/whole split: atom-side percepts and unity/concept evidence.",
            "PartSpace: bottom-up synthesis tower. It owns Sigma, surface-keyed lexicon/MPHF front ends, and the extent/object codebook used for extension learning.",
            "ConceptualSpace: working conceptual interface. It hosts STM, corpus-callosum mixing, and the symbol-table/re-reference machinery that lets nameless percepts cross into reference-bearing symbols.",
            "WholeSpace: top-down analysis tower. It owns Pi, the unified word/symbol codebook, paired orthographic/semantic rows, and codebook-write grammar operations.",
            "WordSpace: grammar and truth host. LanguageLayer scores copy/reduce routes over STM; Grammar provides rules and signatures; GrammarLayer ops perform logic, mereology, lift/lower, and grammatical transformations.",
            "Attention: a single intent code, produced by the parallel parse, primes all relevant codebooks. It biases recognition and retrieval over PS/SS towers without changing grammar-rule dispatch.",
            "Truth and Taxonomy: TruthLayer checks propositions for consistency before admission. AbsoluteTruth is encoded as symbolic truth; RelativeTruth is stored as a taxonomic relation and reflected geometrically as part/whole structure over symbols.",
            "Training runtime: runBatch/runEpoch implement the compute brick, losses, optimizer step, checkpointing, and optional compiled-step path.",
        ],
    )

    doc.add_heading("3. Dynamic Algorithms", level=1)
    add_figure(
        doc,
        images["dynamic"],
        "Figure 2. Dynamic loops: subsymbolic order refines the towers, symbolic order routes grammar over STM, and inter-sentence prediction feeds a future C-prior.",
    )
    add_numbered(
        doc,
        [
            "Subsymbolic order repeats PS -> CS -> WS refinement. The parallel parse produces an intent representation that primes codebook attention across both meronymic towers.",
            "Symbolic order is the serial/grammatical loop. Each word or symbol enters STM, LanguageLayer scores copy/reduce routes, and GrammarLayer operators transform or write symbolic structure.",
            "Corpus callosum mixing is ConceptualSpace-side: percepts cross nameless and factored. Parallel mode uses the explicit 2N -> N mixer; serial mode folds fusion into STM shift/write behavior.",
            "InterSentenceLayer predicts the next STM end-state/root from the LTM chain, stages a C-prior for the next sentence, and contributes L_inter when the prediction is observed.",
            "Training combines output, reconstruction/infill, truth, intra-sentence, inter-sentence, and regularization terms through ModelLoss/Error before the optimizer step.",
        ],
    )

    doc.add_heading("4. Truth, Intension, and Token Promotion", level=1)
    add_figure(
        doc,
        images["truth"],
        "Figure 3. TruthLayer admits consistent propositions; RelativeTruth populates Taxonomy and part/whole symbol geometry; tokens promote to nearest containing types.",
    )
    para(
        doc,
        "Extension learning asks which observations fall under a word. It is learned by the subsymbolic/codebook system from usage, token contexts, perceptual similarity, and nearest-row recognition."
    )
    para(
        doc,
        "Intension learning asks what a word means by virtue of its relations to other propositions and types. A consistent RelativeTruth such as man < mortal becomes a taxonomic containment relation. Geometrically, the symbol for man is part of the symbol for mortal, meaning the equivalence class of man is a subset of the equivalence class of mortal."
    )
    para(
        doc,
        "Token promotion follows containment. If Socrates is categorized as mortal, the taxonomy and codebook geometry should let the system locate the nearest containing symbolic whole, then promote or bind Socrates to the most specific supported type such as man when the evidence warrants it."
    )

    doc.add_heading("5. Concrete Class Inventory", level=1)
    para(
        doc,
        "The inventory below lists the concrete classes that make up the Model, Space, Layer, grammar, and meronymic surfaces. It is intentionally grouped by module rather than drawn as one unreadable inheritance tree."
    )
    inv = class_inventory()
    for module, classes in inv.items():
        doc.add_heading(module, level=2)
        wrapped = textwrap.fill(", ".join(classes), width=95)
        p = para(doc, wrapped)
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string("1B2430")

    doc.add_heading("6. Algorithm Notes", level=1)
    add_bullets(
        doc,
        [
            "XML-driven construction: config resolves dimensions, codebook modes, training knobs, grammar settings, and symbolic/subsymbolic order before BasicModel builds the spaces.",
            "Masked/infill inference: the live within-sentence path is IR-style masked prediction; sentence-level generation is delegated to discourse prediction rather than AR token stepping inside the body.",
            "Meronymic analysis/synthesis: PartSpace synthesizes atoms into wholes with Sigma; WholeSpace analyzes wholes into parts with Pi; MeronymicFoldAdapter binds the membership kernel at those tower slots.",
            "Grammar dispatch: LanguageLayer treats STM as a slab, scores unary copy and binary reduce candidates, then dispatches GrammarLayer operations such as intersection, union, part, equal, query, lift, lower, tense, aspect, and morphology.",
            "Truth admission: grammar may form candidate propositions, but TruthLayer decides whether they are consistent enough to encode. RelativeTruth writes taxonomic structure; AbsoluteTruth writes symbolic truth activations.",
            "Attention over codebooks: semantic relevance is implemented as row-priming over codebook selection/retrieval. It guides focus in the towers while preserving syntactic rule selection as syntactic.",
            "Inter-sentence learning: the sentence root/end-state is appended to an LTM chain; the predictor estimates the next end-state shape and root and scores the estimate against the next observed sentence.",
        ],
    )

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


def pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "Title": ParagraphStyle(
            "Title",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=24,
            leading=29,
            textColor=colors.HexColor("#0B2545"),
            spaceAfter=8,
        ),
        "Subtitle": ParagraphStyle(
            "Subtitle",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=11,
            leading=15,
            textColor=colors.HexColor(MUTED),
            spaceAfter=14,
        ),
        "H1": ParagraphStyle(
            "H1",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=20,
            textColor=colors.HexColor(BLUE),
            spaceBefore=14,
            spaceAfter=8,
        ),
        "H2": ParagraphStyle(
            "H2",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=15,
            textColor=colors.HexColor(DARK_BLUE),
            spaceBefore=10,
            spaceAfter=5,
        ),
        "Body": ParagraphStyle(
            "Body",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=10,
            leading=13.5,
            textColor=colors.HexColor(INK),
            spaceAfter=6,
        ),
        "Small": ParagraphStyle(
            "Small",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=7.8,
            leading=10.2,
            textColor=colors.HexColor(INK),
            spaceAfter=5,
        ),
        "Caption": ParagraphStyle(
            "Caption",
            parent=base["BodyText"],
            fontName="Helvetica-Oblique",
            fontSize=8,
            leading=10,
            alignment=TA_CENTER,
            textColor=colors.HexColor(MUTED),
            spaceBefore=3,
            spaceAfter=10,
        ),
    }


def pdf_p(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(escape(text), style)


def pdf_bullets(items: list[str], style: ParagraphStyle, numbered: bool = False) -> ListFlowable:
    flowables = [ListItem(pdf_p(item, style), leftIndent=12) for item in items]
    return ListFlowable(
        flowables,
        bulletType="1" if numbered else "bullet",
        start="1",
        leftIndent=18,
        bulletFontName="Helvetica",
        bulletFontSize=8,
    )


def pdf_figure(image_path: Path, caption: str, styles: dict[str, ParagraphStyle]) -> list:
    with Image.open(image_path) as img:
        width_px, height_px = img.size
    width = 6.5 * inch
    height = width * height_px / width_px
    return [
        PdfImage(str(image_path), width=width, height=height),
        pdf_p(caption, styles["Caption"]),
    ]


def pdf_footer(canvas_obj, doc_obj) -> None:
    canvas_obj.saveState()
    canvas_obj.setFont("Helvetica", 8)
    canvas_obj.setFillColor(colors.HexColor(MUTED))
    canvas_obj.drawRightString(7.5 * inch, 0.5 * inch, f"BasicModel architecture decomposition - {doc_obj.page}")
    canvas_obj.restoreState()


def build_pdf(images: dict[str, Path]) -> None:
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=1 * inch,
        leftMargin=1 * inch,
        topMargin=0.85 * inch,
        bottomMargin=0.8 * inch,
    )
    story: list = []
    story.append(pdf_p("BasicModel Architecture Decomposition", styles["Title"]))
    story.append(
        pdf_p(
            "Static modules, meronymic towers, grammar/truth machinery, and dynamic learning loops",
            styles["Subtitle"],
        )
    )
    story.append(
        pdf_p(
            "This brief decomposes BasicModel into roughly two dozen architectural blocks, "
            "then expands those blocks into the concrete Model, Space, Layer, grammar, truth, "
            "codebook, and taxonomy modules present in the codebase. It distinguishes extension "
            "learning over codebooks from intension learning through TruthLayer and Taxonomy.",
            styles["Body"],
        )
    )

    story.append(pdf_p("1. Static Architecture", styles["H1"]))
    story.extend(
        pdf_figure(
            images["static"],
            "Figure 1. Static decomposition: entrypoints, model factory/runtime, five-space body, WordSpace grammar/truth host, codebook towers, and training runtime.",
            styles,
        )
    )
    story.append(
        pdf_p(
            "Pipeline. InputSpace -> PartSpace -> ConceptualSpace -> WholeSpace -> OutputSpace. "
            "The WordSpace/WordSubSpace sidecar owns the grammar CPU, semantic attention, truth "
            "admission, taxonomy, and discourse prediction that act on the STM and codebook towers.",
            styles["Body"],
        )
    )

    story.append(pdf_p("2. Modules and Responsibilities", styles["H1"]))
    story.append(
        pdf_bullets(
            [
                "Models: BaseModel owns configuration hydration, training, persistence, and inference scaffolding; BasicModel assembles the space pipeline; ModelFactory validates XML and drives train/eval runs.",
                "InputSpace: converts raw input into the dual view used by the part/whole split: atom-side percepts and unity/concept evidence.",
                "PartSpace: bottom-up synthesis tower. It owns Sigma, surface-keyed lexicon/MPHF front ends, and the extent/object codebook used for extension learning.",
                "ConceptualSpace: working conceptual interface. It hosts STM, corpus-callosum mixing, and the symbol-table/re-reference machinery that lets nameless percepts cross into reference-bearing symbols.",
                "WholeSpace: top-down analysis tower. It owns Pi, the unified word/symbol codebook, paired orthographic/semantic rows, and codebook-write grammar operations.",
                "WordSpace: grammar and truth host. LanguageLayer scores copy/reduce routes over STM; Grammar provides rules and signatures; GrammarLayer ops perform logic, mereology, lift/lower, and grammatical transformations.",
                "Attention: a single intent code, produced by the parallel parse, primes all relevant codebooks. It biases recognition and retrieval over PS/SS towers without changing grammar-rule dispatch.",
                "Truth and Taxonomy: TruthLayer checks propositions for consistency before admission. AbsoluteTruth is encoded as symbolic truth; RelativeTruth is stored as a taxonomic relation and reflected geometrically as part/whole structure over symbols.",
                "Training runtime: runBatch/runEpoch implement the compute brick, losses, optimizer step, checkpointing, and optional compiled-step path.",
            ],
            styles["Body"],
        )
    )

    story.append(PageBreak())
    story.append(pdf_p("3. Dynamic Algorithms", styles["H1"]))
    story.extend(
        pdf_figure(
            images["dynamic"],
            "Figure 2. Dynamic loops: subsymbolic order refines the towers, symbolic order routes grammar over STM, and inter-sentence prediction feeds a future C-prior.",
            styles,
        )
    )
    story.append(
        pdf_bullets(
            [
                "Subsymbolic order repeats PS -> CS -> WS refinement. The parallel parse produces an intent representation that primes codebook attention across both meronymic towers.",
                "Symbolic order is the serial/grammatical loop. Each word or symbol enters STM, LanguageLayer scores copy/reduce routes, and GrammarLayer operators transform or write symbolic structure.",
                "Corpus callosum mixing is ConceptualSpace-side: percepts cross nameless and factored. Parallel mode uses the explicit 2N -> N mixer; serial mode folds fusion into STM shift/write behavior.",
                "InterSentenceLayer predicts the next STM end-state/root from the LTM chain, stages a C-prior for the next sentence, and contributes L_inter when the prediction is observed.",
                "Training combines output, reconstruction/infill, truth, intra-sentence, inter-sentence, and regularization terms through ModelLoss/Error before the optimizer step.",
            ],
            styles["Body"],
            numbered=True,
        )
    )

    story.append(pdf_p("4. Truth, Intension, and Token Promotion", styles["H1"]))
    story.extend(
        pdf_figure(
            images["truth"],
            "Figure 3. TruthLayer admits consistent propositions; RelativeTruth populates Taxonomy and part/whole symbol geometry; tokens promote to nearest containing types.",
            styles,
        )
    )
    for text in [
        "Extension learning asks which observations fall under a word. It is learned by the subsymbolic/codebook system from usage, token contexts, perceptual similarity, and nearest-row recognition.",
        "Intension learning asks what a word means by virtue of its relations to other propositions and types. A consistent RelativeTruth such as man < mortal becomes a taxonomic containment relation. Geometrically, the symbol for man is part of the symbol for mortal, meaning the equivalence class of man is a subset of the equivalence class of mortal.",
        "Token promotion follows containment. If Socrates is categorized as mortal, the taxonomy and codebook geometry should let the system locate the nearest containing symbolic whole, then promote or bind Socrates to the most specific supported type such as man when the evidence warrants it.",
    ]:
        story.append(pdf_p(text, styles["Body"]))

    story.append(PageBreak())
    story.append(pdf_p("5. Concrete Class Inventory", styles["H1"]))
    story.append(
        pdf_p(
            "The inventory below lists the concrete classes that make up the Model, Space, Layer, grammar, and meronymic surfaces. It is grouped by module rather than drawn as one unreadable inheritance tree.",
            styles["Body"],
        )
    )
    for module, classes in class_inventory().items():
        story.append(pdf_p(module, styles["H2"]))
        story.append(pdf_p(", ".join(classes), styles["Small"]))

    story.append(pdf_p("6. Algorithm Notes", styles["H1"]))
    story.append(
        pdf_bullets(
            [
                "XML-driven construction resolves dimensions, codebook modes, training knobs, grammar settings, and symbolic/subsymbolic order before BasicModel builds the spaces.",
                "Masked/infill inference is the live within-sentence path; sentence-level generation is delegated to discourse prediction rather than AR token stepping inside the body.",
                "Meronymic analysis/synthesis: PartSpace synthesizes atoms into wholes with Sigma; WholeSpace analyzes wholes into parts with Pi; MeronymicFoldAdapter binds the membership kernel at those tower slots.",
                "Grammar dispatch treats STM as a slab, scores unary copy and binary reduce candidates, then dispatches GrammarLayer operations such as intersection, union, part, equal, query, lift, lower, tense, aspect, and morphology.",
                "Truth admission separates syntactic proposition formation from admissible structure. RelativeTruth writes taxonomic structure; AbsoluteTruth writes symbolic truth activations.",
                "Attention over codebooks is row-priming over codebook selection/retrieval. It guides focus in the towers while preserving syntactic rule selection as syntactic.",
                "Inter-sentence learning appends the sentence root/end-state to an LTM chain, predicts the next end-state shape and root, and scores the estimate against the next observed sentence.",
            ],
            styles["Body"],
        )
    )

    PDF_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.build(story, onFirstPage=pdf_footer, onLaterPages=pdf_footer)


def main() -> None:
    images = {
        "static": static_architecture_diagram(),
        "dynamic": dynamic_loops_diagram(),
        "truth": truth_taxonomy_diagram(),
    }
    build_docx(images)
    build_pdf(images)
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    main()
